import json
import os
import requests
from docx import Document
from docx.shared import Mm
from docx.shared import Pt
from docx.oxml.ns import qn

header={"User-Agent":"Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/70.0.3538.110 Safari/537.36",
       "Cookie":'输入翻译文档后界面的cookie，详见readme文件'
       }
#url是翻译文档后网页获取的json数据
url = "http://fanyi.youdao.com/trandoc/doc/viewpage?doc=408B55DFA2F84335BC85948146CBB270&client=docserver&keyfrom=doctran&page="

document =  Document()
document.styles['Normal'].font.name = '宋体'
document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
document.styles['Normal'].font.size = Pt(12)
pic_index = 0
dir_path = "fanyi"
file_name = "result.docx" 

# print(os.listdir())
if dir_path not in  os.listdir("."):
    os.mkdir(dir_path)


for i in range(1,999):
    html = requests.get(url + str(i), headers=header)
    data_json = json.loads(html.text)
    print(html.text)
    
    if data_json["errorcode"] != 0:
        print("已经到结尾,完成！！")
        document.save(os.path.join(dir_path,file_name))

        exit(0)
    print("正在处理第"+str(i)+"个JSON文件")
    body = data_json["body"]
    p = 0
    for cell in body:
        s = ''
        for content in cell["trans"]:
            if "tran" not in dict(content).keys():
                pic_url = content["r"][0]["pic"][0]["url"]["val"]
                pic_size_x = content["r"][0]["pic"][0]["spPr"]["xfrm"]["ext"]["cx"]
                pic_size_y = content["r"][0]["pic"][0]["spPr"]["xfrm"]["ext"]["cy"]
                pic = requests.get(pic_url).content
                pic_name = str(pic_index)+".jpg"
                with open(os.path.join(dir_path,pic_name) ,"wb") as f:
                    f.write(pic)
                #图片大小根据需要自行调整，这块优点乱
                document.add_picture(os.path.join(dir_path,pic_name),width=Mm(pic_size_x*3),height=Mm(pic_size_y*3))
                pic_index+=1
            else:
                s +=content["tran"].replace("&nbsp;","")
        document.add_paragraph(s,style="Normal")
        p+=1




















